"""
ChainGuard AI — Images & Diagrams Collection (.docx)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ASSETS_DIR = os.path.join(os.path.dirname(__file__), '_report_assets')

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(26, 26, 46)
    return h

def add_image_with_caption(doc, filename, caption):
    path = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)
            r.italic = True
        doc.add_paragraph() # Spacer
    else:
        doc.add_paragraph(f"[MISSING IMAGE: {filename}]", style='Quote')

def build_collection():
    doc = Document()
    
    # Title
    for _ in range(3): doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ChainGuard AI')
    run.font.size = Pt(36); run.bold = True; run.font.color.rgb = RGBColor(26, 26, 46)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Complete Images & Diagrams Collection')
    run.font.size = Pt(18); run.font.color.rgb = RGBColor(100, 100, 140)

    doc.add_page_break()

    # Section 1: Technical Diagrams
    heading(doc, '1. Technical & Architectural Diagrams')
    
    diagrams = [
        ('system_architecture.png', 'Figure 1: Overall System Architecture (3-Tier Design)'),
        ('agent_workflow.png', 'Figure 2: Multi-Agent Orchestration Flow (LangGraph StateGraph)'),
        ('tech_stack.png', 'Figure 3: Technology Stack Overview'),
        ('ml_pipeline.png', 'Figure 4: Machine Learning Risk Prediction Pipeline'),
        ('er_diagram.png', 'Figure 5: Database Entity-Relationship (ER) Diagram'),
        ('component_tree.png', 'Figure 6: Frontend Component Architecture'),
        ('deployment.png', 'Figure 7: Deployment Architecture (Docker Compose)'),
    ]

    for img, cap in diagrams:
        add_image_with_caption(doc, img, cap)
        if img != diagrams[-1][0]: doc.add_page_break()

    doc.add_page_break()

    # Section 2: Application Screenshots
    heading(doc, '2. Application Interface Screenshots')

    screenshots = [
        ('ss_dashboard.png', 'Figure 8: Main Dashboard - Key Indicators and World Map Viewer'),
        ('ss_events.png', 'Figure 9: Events Monitor - Real-time Global Disruption Feed'),
        ('ss_predictions.png', 'Figure 10: Risk Predictions - 30-Day Forecast & Category Trends'),
        ('ss_mitigation.png', 'Figure 11: Mitigation Strategies - AI-Recommended Action Plans'),
        ('ss_alerts.png', 'Figure 12: Alerts Management - Notification History & Distribution'),
    ]

    for img, cap in screenshots:
        add_image_with_caption(doc, img, cap)
        if img != screenshots[-1][0]: doc.add_page_break()

    # Save
    out = os.path.join(os.path.dirname(__file__), 'ChainGuard_AI_Images_Collection.docx')
    doc.save(out)
    print(f"✅ Images collection saved to: {out}")

if __name__ == '__main__':
    build_collection()
