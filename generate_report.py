"""Generate complete .docx report following End Semester Project Report Guidelines."""
import os, sys
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from report_content import *

OUT = os.path.join(os.path.dirname(__file__), '_report_assets')
os.makedirs(OUT, exist_ok=True)

def shade(cell, c):
    s = cell._element.get_or_add_tcPr()
    s.append(s.makeelement(qn('w:shd'), {qn('w:fill'): c, qn('w:val'): 'clear'}))

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(255,255,255)
        shade(c, '1a1a2e')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 0: shade(c, 'f0f0f5')

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(26, 26, 46)

def img(doc, name, caption):
    doc.add_picture(os.path.join(OUT, name), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(9); r.italic = True

# ── DIAGRAMS ─────────────────────────────────────────────
def box(ax, x, y, w, h, label, color='#1a1a2e', fc='#e8eaf6', fs=9):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.15",fc=fc,ec=color,lw=1.5))
    ax.text(x+w/2,y+h/2,label,ha='center',va='center',fontsize=fs,fontweight='bold',color=color)

def arr(ax, x1,y1,x2,y2):
    ax.annotate('',xy=(x2,y2),xytext=(x1,y1),arrowprops=dict(arrowstyle='->',color='#333',lw=1.5))

def draw_system_arch():
    fig,ax=plt.subplots(figsize=(14,9));ax.set_xlim(0,14);ax.set_ylim(0,9);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(7,8.6,'ChainGuard AI — System Architecture',ha='center',fontsize=16,fontweight='bold',color='#1a1a2e')
    box(ax,0.3,6.5,3.5,1.5,'Frontend Layer\n(React 19 + Vite)\n10 Pages / 6 Components',fc='#bbdefb')
    box(ax,5,6.5,4,1.5,'API Gateway\n(FastAPI + CORS + WebSocket)',fc='#c8e6c9')
    box(ax,10.2,6.5,3.5,1.5,'Hybrid AI Layer\nOllama (Mistral) / GPT-4o',fc='#fff9c4')
    box(ax,2,4,10,2,'',fc='#ede7f6')
    ax.text(7,5.7,'Multi-Agent Orchestration (LangGraph StateGraph)',ha='center',fontsize=11,fontweight='bold',color='#4a148c')
    for i,(ag,c) in enumerate(zip(['Planner\nAgent','Data\nAgent','RAG\nAgent','Risk\nAgent','Decision\nAgent'],['#e3f2fd','#e8f5e9','#fff3e0','#fce4ec','#f3e5f5'])):
        box(ax,2.3+i*1.95,4.2,1.7,1.2,ag,fc=c,fs=8)
    box(ax,0.3,1.5,3,1.5,'Database Layer\nSQLite / PostgreSQL\n(pgvector)',fc='#dcedc8')
    box(ax,5.2,1.5,3.5,1.5,'ML Engine\nXGBoost Classifier\n(risk_model.pkl)',fc='#ffccbc')
    box(ax,10.2,1.5,3.5,1.5,'Memory & RAG\nFAISS / ChromaDB\nVector Store',fc='#d1c4e9')
    arr(ax,3.8,7.25,5,7.25);arr(ax,9,7.25,10.2,7.25);arr(ax,7,6.5,7,6);arr(ax,3.2,4,1.8,3);arr(ax,7,4,7,3);arr(ax,10.5,4,12,3)
    fig.savefig(os.path.join(OUT,'system_architecture.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_agent_workflow():
    fig,ax=plt.subplots(figsize=(13,6));ax.set_xlim(0,13);ax.set_ylim(0,6);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(6.5,5.6,'Multi-Agent Orchestration Flow (LangGraph)',ha='center',fontsize=14,fontweight='bold',color='#1a1a2e')
    nodes=[(0.5,2.5,'START','#78909c'),(2.2,2.5,'Planner\nAgent','#42a5f5'),(4.5,2.5,'Data\nAgent','#66bb6a'),(6.8,2.5,'RAG\nAgent','#ffa726'),(9.1,2.5,'Risk\nAgent','#ef5350'),(11.2,2.5,'Decision\nAgent','#ab47bc')]
    for x,y,l,c in nodes:
        ax.add_patch(plt.Circle((x+0.6,y+0.5),0.55,fc=c,ec='#333',lw=1.5,alpha=0.85))
        ax.text(x+0.6,y+0.5,l,ha='center',va='center',fontsize=7.5,fontweight='bold',color='white')
    ax.add_patch(plt.Circle((12.3,1.0),0.35,fc='#333',ec='#333',lw=2))
    ax.text(12.3,1.0,'END',ha='center',va='center',fontsize=8,fontweight='bold',color='white')
    for i in range(len(nodes)-1):
        ax.annotate('',xy=(nodes[i+1][0]+0.05,nodes[i+1][1]+0.5),xytext=(nodes[i][0]+1.15,nodes[i][1]+0.5),arrowprops=dict(arrowstyle='->',color='#333',lw=2))
    ax.annotate('',xy=(12.3,1.35),xytext=(11.8,2.5),arrowprops=dict(arrowstyle='->',color='#333',lw=2))
    ax.annotate('[GENERAL]',xy=(12.3,1.35),xytext=(2.8,4.2),arrowprops=dict(arrowstyle='->',color='#e91e63',lw=1.5,ls='--'),fontsize=8,color='#e91e63',fontweight='bold')
    fig.savefig(os.path.join(OUT,'agent_workflow.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_tech_stack():
    fig,ax=plt.subplots(figsize=(12,7));ax.set_xlim(0,12);ax.set_ylim(0,7);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(6,6.6,'Technology Stack Overview',ha='center',fontsize=15,fontweight='bold',color='#1a1a2e')
    for name,techs,color,y in [('Presentation',['React 19','Vite','Recharts','Framer Motion','Lucide'],'#bbdefb',5.2),
        ('Application',['FastAPI','LangGraph','LangChain','WebSocket','Pydantic'],'#c8e6c9',3.8),
        ('Intelligence',['XGBoost','Ollama/Mistral','GPT-4o','FAISS/Chroma','scikit-learn'],'#fff9c4',2.4),
        ('Data',['SQLite','PostgreSQL','pgvector','Redis','SQLAlchemy'],'#ffccbc',1.0)]:
        ax.add_patch(FancyBboxPatch((0.5,y),11,1.1,boxstyle="round,pad=0.1",fc=color,ec='#555',lw=1.5))
        ax.text(0.8,y+0.85,name,fontsize=10,fontweight='bold',color='#1a1a2e')
        for i,t in enumerate(techs):
            ax.text(2.8+i*1.8,y+0.35,t,ha='center',fontsize=8,color='#333',bbox=dict(boxstyle='round,pad=0.3',fc='white',ec='#aaa',alpha=0.9))
    fig.savefig(os.path.join(OUT,'tech_stack.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_ml_pipeline():
    fig,ax=plt.subplots(figsize=(12,5));ax.set_xlim(0,12);ax.set_ylim(0,5);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(6,4.6,'ML Risk Prediction Pipeline',ha='center',fontsize=14,fontweight='bold',color='#1a1a2e')
    steps=[(0.5,2,'Training Data\nGeneration\n(5000+ samples)','#e3f2fd'),(3,2,'Feature\nEngineering\n6 Features','#e8f5e9'),(5.5,2,'XGBoost\nClassifier\n(150 trees)','#fff3e0'),(8,2,'Risk\nPrediction\n(probability)','#fce4ec'),(10.5,2,'Risk Level\nClassification\n(4 levels)','#f3e5f5')]
    for x,y,l,c in steps:
        ax.add_patch(FancyBboxPatch((x,y),2,1.8,boxstyle="round,pad=0.15",fc=c,ec='#333',lw=1.5))
        ax.text(x+1,y+0.9,l,ha='center',va='center',fontsize=8,fontweight='bold',color='#1a1a2e')
    for i in range(len(steps)-1):
        ax.annotate('',xy=(steps[i+1][0],2.9),xytext=(steps[i][0]+2,2.9),arrowprops=dict(arrowstyle='->',color='#333',lw=2))
    for i,f in enumerate(['severity','type','weather_severity','distance_to_port','supplier_dependency','historical_freq']):
        ax.text(3.5+(i%3)*1.5,0.8-(i//3)*0.4,f'• {f}',fontsize=7,color='#555')
    fig.savefig(os.path.join(OUT,'ml_pipeline.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_er_diagram():
    fig,ax=plt.subplots(figsize=(13,7));ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(6.5,6.6,'Database Entity-Relationship Diagram',ha='center',fontsize=14,fontweight='bold',color='#1a1a2e')
    for x,y,name,fields,color in [(0.5,3.5,'Event',['id (PK)','type','description','location','severity','timestamp','source'],'#bbdefb'),
        (4.5,4.5,'Supplier',['id (PK)','name','location','reliability_score','tier'],'#c8e6c9'),
        (4.5,1,'Route',['id (PK)','origin','destination','risk_level','distance_km'],'#fff9c4'),
        (8.5,4.5,'RiskPrediction',['id (PK)','event_id (FK)','probability','risk_level','model_version'],'#fce4ec'),
        (8.5,1,'MitigationStrategy',['id (PK)','prediction_id (FK)','title','description','cost'],'#f3e5f5'),
        (0.5,0.5,'Feedback',['id (PK)','prediction_id (FK)','rating','comment','created_at'],'#dcedc8')]:
        h=0.3+len(fields)*0.28
        ax.add_patch(FancyBboxPatch((x,y),3.2,h,boxstyle="round,pad=0.1",fc=color,ec='#333',lw=1.5))
        ax.text(x+1.6,y+h-0.2,name,ha='center',fontsize=9,fontweight='bold',color='#1a1a2e')
        for i,f in enumerate(fields): ax.text(x+0.2,y+h-0.5-i*0.28,f,fontsize=6.5,color='#333')
    ax.plot([3.7,4.5],[4.8,5.0],color='#333',lw=1.5);ax.plot([3.7,4.5],[4.2,2.5],color='#333',lw=1.5)
    ax.plot([7.7,8.5],[5.0,5.0],color='#333',lw=1.5);ax.plot([7.7,8.5],[2.0,2.0],color='#333',lw=1.5,ls='--')
    fig.savefig(os.path.join(OUT,'er_diagram.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_component_tree():
    fig,ax=plt.subplots(figsize=(14,8));ax.set_xlim(0,14);ax.set_ylim(0,8);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(7,7.6,'Frontend Component Architecture (5-Phase)',ha='center',fontsize=14,fontweight='bold',color='#1a1a2e')
    ax.add_patch(FancyBboxPatch((5.5,6.5),3,0.7,boxstyle="round,pad=0.1",fc='#1a1a2e',ec='#333',lw=2))
    ax.text(7,6.85,'App.jsx (Router)',ha='center',fontsize=10,fontweight='bold',color='white')
    for x,l in [(1,'Sidebar'),(5.5,'TopBar'),(10,'ChatbotWidget')]:
        ax.add_patch(FancyBboxPatch((x,5.3),2.5,0.7,boxstyle="round,pad=0.1",fc='#78909c',ec='#333',lw=1.5))
        ax.text(x+1.25,5.65,l,ha='center',fontsize=8,fontweight='bold',color='white')
        ax.plot([7,x+1.25],[6.5,6.0],color='#555',lw=1)
    for name,pages,color,x in [('Phase 1',['DataIngestion','EventsMonitor'],'#e3f2fd',0.3),('Phase 2',['KnowledgeBase'],'#e8f5e9',3.3),
        ('Phase 3',['AgentBrain','Predictions'],'#fff3e0',5.5),('Phase 4',['Mitigation'],'#fce4ec',8.5),('Phase 5',['Dashboard','Alerts','Reports','Feedback'],'#f3e5f5',10.8)]:
        h=0.5+len(pages)*0.4
        ax.add_patch(FancyBboxPatch((x,4.2-h),2.5,h,boxstyle="round,pad=0.1",fc=color,ec='#555',lw=1))
        ax.text(x+1.25,4.1,name,ha='center',fontsize=7,fontweight='bold',color='#1a1a2e')
        for i,p in enumerate(pages): ax.text(x+0.3,3.6-i*0.38,f'• {p}.jsx',fontsize=6.5,color='#333')
        ax.plot([7,x+1.25],[6.5,4.2],color='#aaa',lw=0.8,ls='--')
    fig.savefig(os.path.join(OUT,'component_tree.png'),dpi=180,bbox_inches='tight');plt.close()

def draw_deployment():
    fig,ax=plt.subplots(figsize=(12,6));ax.set_xlim(0,12);ax.set_ylim(0,6);ax.axis('off');fig.patch.set_facecolor('#f8f9fa')
    ax.text(6,5.6,'Deployment Architecture (Docker Compose)',ha='center',fontsize=14,fontweight='bold',color='#1a1a2e')
    ax.add_patch(FancyBboxPatch((0.2,1.8),11.6,3.2,boxstyle="round,pad=0.2",fc='none',ec='#2196f3',lw=2,ls='--'))
    ax.text(6,1.3,'Docker Compose Network',ha='center',fontsize=10,color='#2196f3',fontweight='bold')
    containers=[(0.5,2.5,'Frontend\n(Nginx+React)\n:5173','#bbdefb'),(3.5,2.5,'Backend\n(FastAPI)\n:8000','#c8e6c9'),(6.5,2.5,'PostgreSQL\n(pgvector)\n:5432','#fff9c4'),(9.5,2.5,'Redis\n(Cache)\n:6379','#ffccbc')]
    for x,y,l,c in containers:
        ax.add_patch(FancyBboxPatch((x,y),2.5,2,boxstyle="round,pad=0.15",fc=c,ec='#333',lw=1.5))
        ax.text(x+1.25,y+1,l,ha='center',va='center',fontsize=8,fontweight='bold',color='#1a1a2e')
    for i in range(len(containers)-1):
        ax.annotate('',xy=(containers[i+1][0],3.5),xytext=(containers[i][0]+2.5,3.5),arrowprops=dict(arrowstyle='<->',color='#333',lw=1.5))
    ax.add_patch(FancyBboxPatch((4,0.2),4,0.7,boxstyle="round,pad=0.1",fc='#e1bee7',ec='#333',lw=1))
    ax.text(6,0.55,'Ollama (Local AI) — localhost:11434',ha='center',fontsize=8,fontweight='bold',color='#333')
    ax.annotate('',xy=(6,1.8),xytext=(6,0.9),arrowprops=dict(arrowstyle='->',color='#9c27b0',lw=1.5,ls='--'))
    fig.savefig(os.path.join(OUT,'deployment.png'),dpi=180,bbox_inches='tight');plt.close()

def gen_diagrams():
    print("Generating 7 diagrams...")
    draw_system_arch(); draw_agent_workflow(); draw_tech_stack()
    draw_ml_pipeline(); draw_er_diagram(); draw_component_tree(); draw_deployment()
    print("Done.")

# ── BUILD DOCX ───────────────────────────────────────────
def build():
    gen_diagrams()
    doc = Document()
    for s in doc.sections: s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)

    # 1. TITLE PAGE
    for _ in range(4): doc.add_paragraph()
    for text,sz,clr in [(TITLE_INFO['title'],36,(26,26,46)),(TITLE_INFO['subtitle'],18,(100,100,140)),
        (TITLE_INFO['domain'],14,(80,80,120))]:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(text);r.font.size=Pt(sz);r.bold=(sz==36);r.font.color.rgb=RGBColor(*clr)
    doc.add_paragraph()
    for line in [f"Student: {TITLE_INFO['student']}",f"Roll Number: {TITLE_INFO['roll']}",
        f"Course: {TITLE_INFO['course']}",f"Instructor: {TITLE_INFO['instructor']}",f"Date: {TITLE_INFO['date']}"]:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(line);r.font.size=Pt(11)
    doc.add_page_break()

    # TOC
    H(doc,'Table of Contents')
    for i,t in enumerate(['Abstract','Problem Definition & Domain','System Architecture & Agent Roles',
        'Memory Integration & Knowledge Retrieval (RAG)','Tool Usage & Environment Setup',
        'Agent Planning & Orchestration','Validation and Testing of Agentic Systems',
        'System Evaluation','Results and Insights','Conclusion','Limitations','Future Work','References'],1):
        doc.add_paragraph(f'{i}. {t}',style='List Number')
    doc.add_page_break()

    # 2. ABSTRACT
    H(doc,'1. Abstract'); doc.add_paragraph(ABSTRACT); doc.add_page_break()

    # 3. PROBLEM DEFINITION & DOMAIN
    H(doc,'2. Problem Definition & Domain')
    for title,text in PROBLEM_DEFINITION:
        H(doc,title,level=2); doc.add_paragraph(text)
    doc.add_page_break()

    # 4. SYSTEM ARCHITECTURE & AGENT ROLES
    H(doc,'3. System Architecture & Agent Roles')
    doc.add_paragraph('The system follows a hybrid AI architecture combining cloud and local intelligence across three tiers.')
    img(doc,'system_architecture.png','Figure 1: ChainGuard AI System Architecture')
    doc.add_paragraph()
    add_table(doc,['Agent','Role','Model Used','Online/Offline'],[
        ['Planner Agent','Decomposes queries, routes workflow','Mistral-7B / GPT-4o','Both'],
        ['Data Collection Agent','Fetches real-time events from DB','None (SQLAlchemy)','Offline'],
        ['RAG Agent','Retrieves historical patterns','FAISS/ChromaDB','Offline'],
        ['Risk Analysis Agent','Predicts disruption probability','XGBoost (risk_model.pkl)','Offline'],
        ['Decision Agent','Generates mitigation strategies','Mistral-7B / GPT-4o','Both'],
    ])
    doc.add_paragraph()
    img(doc,'agent_workflow.png','Figure 2: Multi-Agent Orchestration Flow (LangGraph)')
    doc.add_page_break()

    # 5. MEMORY INTEGRATION & RAG
    H(doc,'4. Memory Integration & Knowledge Retrieval (RAG)')
    for title,text in MEMORY_RAG:
        H(doc,title,level=2); doc.add_paragraph(text)
    doc.add_page_break()

    # 6. TOOL USAGE & ENVIRONMENT SETUP
    H(doc,'5. Tool Usage & Environment Setup')
    H(doc,'5.1 Tools List',level=2)
    add_table(doc,['Tool','Purpose'],TOOL_USAGE[0][1])
    doc.add_paragraph()
    H(doc,'5.2 Agent-Tool Interface',level=2); doc.add_paragraph(TOOL_USAGE[1][1])
    H(doc,'5.3 Example Inputs & Outputs',level=2); doc.add_paragraph(TOOL_USAGE[2][1])
    img(doc,'tech_stack.png','Figure 3: Technology Stack Overview')
    doc.add_page_break()

    # 7. AGENT PLANNING & ORCHESTRATION
    H(doc,'6. Agent Planning & Orchestration')
    for k in ['7.1','7.2','7.3']:
        title,text = PLANNING_ORCHESTRATION[k]
        H(doc,f'6.{k[-1]} {title}',level=2); doc.add_paragraph(text)
    doc.add_page_break()

    # 8. VALIDATION AND TESTING
    H(doc,'7. Validation and Testing of Agentic Systems')
    for key in ['unit_testing','trajectory','consistency','robustness','rag_metrics','hitl']:
        title,text = VALIDATION_TESTING[key]
        H(doc,title,level=2); doc.add_paragraph(text)
        doc.add_paragraph()
    doc.add_page_break()

    # ML Pipeline diagram
    H(doc,'7.1 ML Model Validation',level=2)
    img(doc,'ml_pipeline.png','Figure 4: ML Risk Prediction Pipeline')
    add_table(doc,['Feature','Type','Range','Description'],[
        ['severity','Categorical','1-4','Event severity (low→critical)'],
        ['type','Categorical','1-6','Event type (hurricane→geopolitical)'],
        ['weather_severity','Numeric','0-5','Weather severity index'],
        ['distance_to_port','Numeric','5-1000 km','Proximity to supply chain port'],
        ['supplier_dependency_score','Numeric','10-100%','Supplier dependency level'],
        ['historical_disruption_frequency','Numeric','0-50/yr','Past disruption frequency'],
    ])
    doc.add_paragraph()
    img(doc,'er_diagram.png','Figure 5: Database Entity-Relationship Diagram')
    doc.add_page_break()

    # 9. SYSTEM EVALUATION
    H(doc,'8. System Evaluation'); doc.add_paragraph(SYSTEM_EVALUATION)
    add_table(doc,['Metric','Local (Mistral-7B)','Cloud (GPT-4o-mini)','XGBoost ML'],[
        ['Response Time','~0.5s/call','~1.2s/call','<50ms'],
        ['Quality (Reasoning)','Good (structured)','Excellent (nuanced)','N/A (deterministic)'],
        ['Offline Capable','Yes','No','Yes'],
        ['Cost per Query','Free','~$0.002','Free'],
        ['Consistency','High','Medium (non-deterministic)','100%'],
        ['Context Window','4K-8K tokens','128K tokens','6 features'],
    ])
    doc.add_page_break()

    # 10. RESULTS AND INSIGHTS
    H(doc,'9. Results and Insights')
    for r in RESULTS_INSIGHTS: doc.add_paragraph(r, style='List Bullet')
    img(doc,'component_tree.png','Figure 6: Frontend Component Architecture (5-Phase)')
    doc.add_page_break()

    # 11. CONCLUSION
    H(doc,'10. Conclusion'); doc.add_paragraph(CONCLUSION)

    # 12. LIMITATIONS
    H(doc,'11. Limitations')
    for l in LIMITATIONS: doc.add_paragraph(l, style='List Bullet')

    # 13. FUTURE WORK
    H(doc,'12. Future Work')
    for e in ['Integration with real-time news/weather APIs (NewsAPI, OpenWeatherMap)',
        'Advanced RAG with production vector database (Pinecone/Weaviate)',
        'Model retraining pipeline with MLflow tracking','Multi-language support for global deployment',
        'Mobile-responsive PWA version','Role-based access control (RBAC)',
        'Anomaly detection using time-series models (Prophet/LSTM)','Supply chain digital twin simulation']:
        doc.add_paragraph(e, style='List Bullet')

    # 14. REFERENCES
    H(doc,'13. References')
    for i,ref in enumerate(['LangGraph Documentation — https://langchain-ai.github.io/langgraph/',
        'FastAPI Framework — https://fastapi.tiangolo.com/','XGBoost — https://xgboost.readthedocs.io/',
        'React 19 — https://react.dev/','Ollama — https://ollama.ai/',
        'Docker Compose — https://docs.docker.com/compose/','Recharts — https://recharts.org/',
        'SQLAlchemy Async — https://docs.sqlalchemy.org/'],1):
        doc.add_paragraph(f'[{i}] {ref}')

    # Deployment diagram at end
    img(doc,'deployment.png','Figure 7: Deployment Architecture (Docker Compose)')

    out = os.path.join(os.path.dirname(__file__), 'ChainGuard_AI_Agentic_Report_v2.docx')
    doc.save(out)
    print(f"\n✅ Report saved: {out}")

if __name__ == '__main__':
    build()
